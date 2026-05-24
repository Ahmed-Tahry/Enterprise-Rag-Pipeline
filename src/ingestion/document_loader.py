"""
src/ingestion/document_loader.py
=================================
Multi-format document loader supporting PDF, DOCX, HTML, TXT, and Markdown.

Each loader returns a list of Document objects with:
  - page_content: raw extracted text
  - metadata: source path, page number, file type, title, etc.

Real-world insight: 80% of enterprise data lives in PDFs and Word docs.
Getting clean text extraction right (handling tables, headers, footers,
multi-column layouts) is where most RAG projects fail in production.
"""

import os
import re
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Optional, Dict, Any
from loguru import logger


@dataclass
class Document:
    """A loaded document chunk with text and metadata."""
    page_content: str
    metadata: Dict[str, Any] = field(default_factory=dict)

    def __repr__(self):
        preview = self.page_content[:80].replace("\n", " ")
        return f"Document(source='{self.metadata.get('source', '?')}', preview='{preview}...')"


class PDFLoader:
    """
    Load PDF files using pypdf.
    
    Handles:
    - Multi-page PDFs
    - Extracts page numbers as metadata
    - Cleans common PDF artifacts (ligatures, hyphenation, extra whitespace)
    """
    
    def load(self, file_path: str) -> List[Document]:
        try:
            import pypdf
        except ImportError:
            raise ImportError("Install pypdf: pip install pypdf")
        
        path = Path(file_path)
        documents = []
        
        with open(file_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            total_pages = len(reader.pages)
            
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                text = self._clean_pdf_text(text)
                
                if text.strip():
                    documents.append(Document(
                        page_content=text,
                        metadata={
                            "source": str(path),
                            "filename": path.name,
                            "file_type": "pdf",
                            "page": page_num + 1,
                            "total_pages": total_pages,
                        }
                    ))
        
        logger.info(f"Loaded PDF: {path.name} ({total_pages} pages, {len(documents)} non-empty)")
        return documents
    
    def _clean_pdf_text(self, text: str) -> str:
        """Remove common PDF extraction artifacts."""
        # Fix ligatures (ﬁ → fi, ﬂ → fl)
        ligatures = {"ﬁ": "fi", "ﬂ": "fl", "ﬀ": "ff", "ﬃ": "ffi", "ﬄ": "ffl"}
        for lig, rep in ligatures.items():
            text = text.replace(lig, rep)
        # Remove hyphenation at line breaks
        text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
        # Normalise whitespace
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()


class DocxLoader:
    """
    Load Microsoft Word (.docx) files.
    Preserves paragraph structure and extracts heading hierarchy.
    """
    
    def load(self, file_path: str) -> List[Document]:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")
        
        path = Path(file_path)
        doc = DocxDocument(file_path)
        
        sections: List[Dict] = []
        current_section = {"heading": None, "paragraphs": []}
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Detect headings
            if para.style.name.startswith("Heading"):
                if current_section["paragraphs"]:
                    sections.append(current_section)
                current_section = {"heading": text, "paragraphs": []}
            else:
                current_section["paragraphs"].append(text)
        
        if current_section["paragraphs"]:
            sections.append(current_section)
        
        documents = []
        for i, section in enumerate(sections):
            content_parts = []
            if section["heading"]:
                content_parts.append(f"## {section['heading']}")
            content_parts.extend(section["paragraphs"])
            content = "\n\n".join(content_parts)
            
            if content.strip():
                documents.append(Document(
                    page_content=content,
                    metadata={
                        "source": str(path),
                        "filename": path.name,
                        "file_type": "docx",
                        "section": i + 1,
                        "heading": section["heading"] or "Untitled Section",
                    }
                ))
        
        logger.info(f"Loaded DOCX: {path.name} ({len(sections)} sections)")
        return documents


class HTMLLoader:
    """
    Load HTML files, stripping tags and preserving structure.
    Useful for loading web pages, documentation sites, wikis.
    """
    
    def load(self, file_path: str) -> List[Document]:
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("Install beautifulsoup4: pip install beautifulsoup4")
        
        path = Path(file_path)
        
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            html = f.read()
        
        soup = BeautifulSoup(html, "html.parser")
        
        # Remove script/style/nav elements
        for tag in soup(["script", "style", "nav", "header", "footer", "aside"]):
            tag.decompose()
        
        title = soup.title.string if soup.title else path.stem
        text = soup.get_text(separator="\n")
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        
        documents = []
        if text:
            documents.append(Document(
                page_content=text,
                metadata={
                    "source": str(path),
                    "filename": path.name,
                    "file_type": "html",
                    "title": str(title),
                }
            ))
        
        logger.info(f"Loaded HTML: {path.name} (title: {title})")
        return documents


class TextLoader:
    """Load plain text and Markdown files."""
    
    def load(self, file_path: str) -> List[Document]:
        path = Path(file_path)
        
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        
        documents = []
        if text.strip():
            documents.append(Document(
                page_content=text.strip(),
                metadata={
                    "source": str(path),
                    "filename": path.name,
                    "file_type": path.suffix.lstrip(".") or "txt",
                }
            ))
        
        logger.info(f"Loaded text: {path.name} ({len(text)} chars)")
        return documents


class DocumentLoader:
    """
    Unified document loader — automatically selects the right loader
    based on file extension.
    
    Usage:
        loader = DocumentLoader()
        docs = loader.load("report.pdf")
        docs = loader.load_directory("./docs/")
    """
    
    LOADERS = {
        ".pdf":  PDFLoader,
        ".docx": DocxLoader,
        ".doc":  DocxLoader,
        ".html": HTMLLoader,
        ".htm":  HTMLLoader,
        ".txt":  TextLoader,
        ".md":   TextLoader,
        ".rst":  TextLoader,
    }
    
    def load(self, file_path: str) -> List[Document]:
        """Load a single file."""
        path = Path(file_path)
        ext = path.suffix.lower()
        
        if ext not in self.LOADERS:
            raise ValueError(f"Unsupported file type: {ext}. Supported: {list(self.LOADERS.keys())}")
        
        loader = self.LOADERS[ext]()
        return loader.load(file_path)
    
    def load_directory(
        self,
        directory: str,
        recursive: bool = True,
        extensions: Optional[List[str]] = None,
    ) -> List[Document]:
        """
        Load all supported documents from a directory.
        
        Args:
            directory:  Path to directory
            recursive:  Search subdirectories
            extensions: Filter to specific extensions (e.g. ['.pdf', '.docx'])
        """
        dir_path = Path(directory)
        if not dir_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")
        
        allowed_exts = extensions or list(self.LOADERS.keys())
        
        pattern = "**/*" if recursive else "*"
        all_files = [
            f for f in dir_path.glob(pattern)
            if f.is_file() and f.suffix.lower() in allowed_exts
        ]
        
        all_documents = []
        failed = []
        
        for file_path in sorted(all_files):
            try:
                docs = self.load(str(file_path))
                all_documents.extend(docs)
            except Exception as e:
                logger.warning(f"Failed to load {file_path.name}: {e}")
                failed.append(str(file_path))
        
        logger.info(
            f"Loaded {len(all_documents)} document pages from {len(all_files)} files "
            f"({len(failed)} failed)"
        )
        return all_documents
